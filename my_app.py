from docx import Document
from docx.shared import Inches
import pyttsx3 


def speak(text):
    pyttsx3.speak(text)

document = Document()

name = input('What is your name?')
speak('Hello '+name+ ',how are you today?')
phone = input('Enter your phone number:')
email = input('enter your email:')

document.add_picture('12FE1A0323.jpg',width=Inches(1.5))
document.add_paragraph(name+' | '+phone+' | '+email)

# About me
document.add_heading('About me')
document.add_paragraph(input("Tell me about yourself ?"))

# Work Experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company_name = input('What is your previous company ?')
start_date = input('From date ?')
end_date = input('To date ?')

p.add_run(company_name+'').bold=True
p.add_run(' '+start_date+' to '+end_date+'\n').italic = True
experience_Details =  input('Enter your experience at company '+company_name)

p.add_run(experience_Details)

# more experiences 

while True:
    has_more_experience = input('Do you have more experience? yes/no :')
    if has_more_experience.lower() == 'yes':
        p = document.add_paragraph()

        company_name = input('What is your previous company ?')
        start_date = input('From date ?')
        end_date = input('To date ?')

        p.add_run(company_name+'').bold=True
        p.add_run(' '+start_date+' to '+end_date+'\n').italic = True
        experience_Details =  input('Enter your experience at company '+company_name)

        p.add_run(experience_Details)
    else:
        break


#skills
document.add_heading('Skills ')
skill =  input('Enter your skill')
p = document.add_paragraph(skill)
p.style='List Bullet'

while True:
    has_more_skills = input("Do you have more skills? yes/no :")
    if has_more_skills.lower() == 'yes':
        skill =  input('Enter your skill')
        p = document.add_paragraph(skill)
        p.style='List Bullet'
    else:
        break


# footer
section = document.sections[0]
footer = section.footer
para = footer.paragraphs[0]
para.text = 'CV has generated using docs api of python'


document.save('Hari_CV.docx')